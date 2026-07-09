#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate binary documents for biz-security knowledge base."""

import os
from pathlib import Path

BASE = Path(__file__).resolve().parent


def generate_docx():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)

    title = doc.add_heading("星云科技集团数据分类分级规范", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("文件编号：XY-SEC-DC-2025-001  |  版本：V1.5  |  密级：内部公开\n")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(80, 80, 80)
    meta.add_run("归口部门：信息安全部  |  批准人：首席信息安全官（CISO）")

    doc.add_paragraph()

    sections = [
        ("1. 目的与范围", [
            "为统一集团数据资产管理口径，建立可执行、可审计、可自动化的数据分类分级体系，"
            "支撑访问控制、加密存储、脱敏展示、日志审计及等保合规要求，特制定本规范。",
            "本规范适用于集团及下属单位所有结构化数据（数据库表/字段）、非结构化数据（文档、"
            "图片、日志）以及系统接口传输中的数据对象。",
        ]),
        ("2. 术语定义", [
            "数据分类：按业务属性划分数据类别，如个人信息、财务数据、知识产权、运营数据等。",
            "数据分级：按泄露后对国家安全、公共利益、个人权益、业务连续性的影响程度，划分为 L1～L4。",
            "数据标签（Data Tag）：附着在字段/文件/接口上的机读标识，供策略引擎自动执行控制策略。",
            "数据 Owner：对特定数据域负有分类分级、授权审批、质量与安全责任的业务负责人。",
        ]),
        ("3. 数据分类体系", [
            "集团采用「一级类别 + 二级子类」二维分类法，共 8 个一级类别：",
        ]),
    ]

    for heading, paras in sections:
        doc.add_heading(heading, level=1)
        for p in paras:
            doc.add_paragraph(p)

    # Classification table
    table = doc.add_table(rows=9, cols=3)
    table.style = "Table Grid"
    headers = ["一级类别", "二级子类", "典型示例"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    rows_data = [
        ("A 组织管理", "制度流程、会议纪要", "员工手册、董事会纪要"),
        ("B 人力资源", "员工档案、薪酬考勤", "劳动合同、工资条、绩效"),
        ("C 财务资金", "核算报表、税务票据", "总账、发票、银行回单"),
        ("D 客户市场", "客户资料、合同商机", "CRM 客户、销售合同"),
        ("E 研发技术", "源代码、设计文档", "Git 仓库、架构图、专利"),
        ("F 运营管理", "日志监控、配置参数", "访问日志、监控指标"),
        ("G 个人信息", "身份联系、行为偏好", "姓名手机身份证、浏览记录"),
        ("H 监管合规", "审计报告、等保材料", "测评报告、监管报送"),
    ]
    for i, row in enumerate(rows_data, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val

    doc.add_heading("4. 数据分级标准", level=1)
    doc.add_paragraph(
        "分级判定采用「就高不就低」原则：同一记录含多个字段时，记录整体等级取最高字段等级。"
    )

    level_table = doc.add_table(rows=5, cols=5)
    level_table.style = "Table Grid"
    lh = ["等级", "名称", "泄露影响", "典型数据", "控制要求摘要"]
    for i, h in enumerate(lh):
        level_table.rows[0].cells[i].text = h
    levels = [
        ("L1", "公开", "无损害或已公开", "产品手册、招聘公告", "基础访问控制"),
        ("L2", "内部", "轻微影响运营", "内部制度、会议纪要", "认证+授权+基础审计"),
        ("L3", "敏感", "较大影响个人/业务", "客户联系方式、合同金额", "字段加密+脱敏+强审计"),
        ("L4", "核心机密", "严重损害或违法", "身份证、银行卡、源码、密钥", "国密/HSM+最小可见+双人审批"),
    ]
    for i, row in enumerate(levels, start=1):
        for j, val in enumerate(row):
            level_table.rows[i].cells[j].text = val

    doc.add_heading("5. 分级判定流程", level=1)
    steps = [
        "业务方梳理数据资产清单，标注业务含义与使用场景；",
        "数据 Owner 按本规范第 4 章初步定级，信息安全部复核；",
        "在元数据平台打标（tag），同步至数据库、API 网关、DLP；",
        "系统变更涉及新字段时，须在需求评审阶段完成分级标注；",
        "每年度进行一次分级复核，业务重大变化时随时更新。",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}", style="List Number")

    doc.add_heading("6. 标签命名规范", level=1)
    doc.add_paragraph("标签格式：XY.DATA.{类别}.{等级}.{子类}")
    doc.add_paragraph("示例：")
    examples = [
        "XY.DATA.G.L4.ID_CARD — 个人身份证号码",
        "XY.DATA.C.L3.INVOICE — 增值税发票信息",
        "XY.DATA.E.L4.SOURCE_CODE — 源代码",
        "XY.DATA.D.L3.CONTRACT — 销售合同正文",
    ]
    for ex in examples:
        doc.add_paragraph(ex, style="List Bullet")

    doc.add_heading("7. 差异化控制策略", level=1)
    ctrl_table = doc.add_table(rows=5, cols=6)
    ctrl_table.style = "Table Grid"
    ch = ["等级", "存储加密", "传输", "展示", "导出", "日志留存"]
    for i, h in enumerate(ch):
        ctrl_table.rows[0].cells[i].text = h
    ctrl = [
        ("L1", "可选", "HTTPS", "明文", "允许", "90天"),
        ("L2", "推荐", "HTTPS+内网", "明文", "审批", "180天"),
        ("L3", "AES-256必须", "TLS1.2+", "脱敏", "总监审批", "365天"),
        ("L4", "国密/HSM", "mTLS/专线", "最小可见+水印", "CISO审批", "365天"),
    ]
    for i, row in enumerate(ctrl, start=1):
        for j, val in enumerate(row):
            ctrl_table.rows[i].cells[j].text = val

    doc.add_heading("8. 个人信息特别要求", level=1)
    doc.add_paragraph(
        "涉及个人信息的数据，无论业务分类如何，最低定级为 L3；"
        "身份证号、生物识别、金融账户等直接识别个人的字段必须定为 L4。"
    )
    doc.add_paragraph("个人信息处理须额外满足：")
    pi_items = [
        "具有明确、合理的处理目的，与处理目的直接相关；",
        "采取对个人权益影响最小的方式；",
        "事先进行个人信息保护影响评估（PIA）；",
        "向个人告知处理规则，取得有效同意（法律豁免情形除外）。",
    ]
    for item in pi_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("9. 违规与问责", level=1)
    doc.add_paragraph(
        "未经审批将高等级数据降级标注、擅自导出 L4 数据、在公网传输未加密 L3 数据等行为，"
        "按《信息安全管理制度》第 12 章处理。"
    )

    doc.add_heading("10. 附则", level=1)
    doc.add_paragraph("本规范由信息安全部负责解释，与《信息安全管理制度》配套使用。")
    doc.add_paragraph("自 2025 年 1 月 1 日起施行。")

    out = BASE / "数据分类分级规范.docx"
    doc.save(out)
    print(f"Created: {out}")


def generate_xlsx():
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "权限申请对照表"

    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill("solid", fgColor="2F5496")
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    headers = [
        "角色编码", "角色名称", "适用系统", "数据最高等级",
        "权限范围", "是否特权", "审批层级", "复核周期", "备注",
    ]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = thin

    data = [
        ("ROLE-OA-USER", "OA普通用户", "SYS-OA-001", "L2", "发起审批、查看本人公文", "否", "直属经理", "季度", "默认入职角色"),
        ("ROLE-OA-HR", "OA人事专员", "SYS-OA-001", "L3", "人事审批、员工档案只读", "否", "HR总监", "季度", "须完成隐私培训"),
        ("ROLE-OA-ADMIN", "OA系统管理员", "SYS-OA-001", "L3", "流程配置、组织同步", "是", "行政总监+IT", "月度", "堡垒机+MFA"),
        ("ROLE-ERP-ACCT", "ERP会计", "SYS-ERP-002", "L3", "凭证录入、报表查询", "否", "财务经理", "季度", ""),
        ("ROLE-ERP-CFO", "ERP财务总监", "SYS-ERP-002", "L4", "全模块、含薪酬总账", "是", "CFO+CISO", "月度", "双人复核支付"),
        ("ROLE-ERP-DBA", "ERP数据库管理员", "SYS-ERP-002", "L4", "DB读写、备份恢复", "是", "CISO", "月度", "仅堡垒机操作"),
        ("ROLE-CRM-SALES", "CRM销售代表", "SYS-CRM-003", "L3", "本人客户线索商机", "否", "销售经理", "季度", "禁止批量导出"),
        ("ROLE-CRM-MGR", "CRM销售总监", "SYS-CRM-003", "L3", "本部门全部客户", "否", "销售VP", "季度", "导出须工单"),
        ("ROLE-HR-BP", "HR业务伙伴", "SYS-HR-004", "L4", "所辖部门薪酬档案", "否", "HR总监+CISO", "月度", "含L4薪酬数据"),
        ("ROLE-HR-ADMIN", "eHR系统管理员", "SYS-HR-004", "L4", "全量员工档案配置", "是", "CISO", "月度", "操作全审计"),
        ("ROLE-DMP-ANALYST", "数据分析师", "SYS-DMP-005", "L3", "指标查询、脱敏数据集", "否", "数据总监", "季度", "禁止下载明细"),
        ("ROLE-DMP-ADMIN", "数据中台管理员", "SYS-DMP-005", "L4", "库表管理、调度配置", "是", "CISO", "月度", "生产只读分离"),
        ("ROLE-DEV-ENG", "研发工程师", "SYS-DEV-006", "L3", "代码读写、CI触发", "否", "研发经理", "季度", "禁止主分支强推"),
        ("ROLE-DEV-LEAD", "研发负责人", "SYS-DEV-006", "L4", "主分支合并、密钥管理", "是", "研发总监", "月度", "密钥不落盘"),
        ("ROLE-SOC-ANALYST", "SOC分析员", "SYS-SOC-010", "L3", "告警研判、日志查询", "否", "安全经理", "季度", "禁止改规则"),
        ("ROLE-SOC-ADMIN", "SOC管理员", "SYS-SOC-010", "L4", "规则配置、封禁策略", "是", "CISO", "月度", "变更双人复核"),
        ("ROLE-IAM-ADMIN", "IAM管理员", "SYS-SSO-009", "L4", "账号创建、角色绑定", "是", "CISO", "月度", "禁止给自己授权"),
        ("ROLE-IT-OPS", "IT运维工程师", "全部生产系统", "L3", "监控、补丁、重启", "是", "运维经理", "月度", "禁止直连接数据库"),
        ("ROLE-IT-OPS-SUPER", "IT运维超级管理员", "全部生产系统", "L4", "root/域管/网络配置", "是", "CISO", "周度", "PAM+录屏"),
        ("ROLE-AUDIT-RO", "内审只读", "全部系统", "L4", "全系统只读审计", "否", "内审总监", "季度", "只读账号专用"),
        ("ROLE-API-PARTNER", "开放平台合作方", "SYS-API-013", "L2", "授权API调用", "否", "业务Owner", "季度", "限流+IP白名单"),
        ("ROLE-INS-OPS", "保险核心运维", "SYS-INS-017", "L4", "保单理赔系统运维", "是", "CISO+合规", "月度", "监管专有云"),
    ]

    for row_idx, row_data in enumerate(data, 2):
        for col_idx, val in enumerate(row_data, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.alignment = Alignment(vertical="center", wrap_text=True)
            cell.border = thin
            if col_idx == 6 and val == "是":
                cell.fill = PatternFill("solid", fgColor="FCE4D6")

    widths = [16, 18, 16, 12, 28, 10, 16, 10, 22]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[chr(64 + i)].width = w

    # Sheet 2: approval matrix
    ws2 = wb.create_sheet("审批层级说明")
    ws2["A1"] = "审批层级"
    ws2["B1"] = "审批人"
    ws2["C1"] = "适用场景"
    ws2["D1"] = "SLA"
    for c in "ABCD":
        ws2[f"{c}1"].font = header_font
        ws2[f"{c}1"].fill = header_fill

    approval = [
        ("直属经理", "员工直属上级", "L1/L2 只读、基础操作", "1工作日"),
        ("部门总监", "一级部门负责人", "L3 读写、跨组数据", "2工作日"),
        ("系统Owner", "业务系统负责人", "生产写权限、批量查询", "2工作日"),
        ("CISO", "首席信息安全官", "L4数据、特权账号、DB写", "3工作日"),
        ("安全委员会", "信息安全委员会", "跨境传输、批量脱敏例外", "5工作日"),
    ]
    for i, row in enumerate(approval, 2):
        for j, val in enumerate(row, 1):
            ws2.cell(row=i, column=j, value=val)

    ws2.column_dimensions["A"].width = 14
    ws2.column_dimensions["B"].width = 20
    ws2.column_dimensions["C"].width = 30
    ws2.column_dimensions["D"].width = 12

    out = BASE / "权限申请对照表.xlsx"
    wb.save(out)
    print(f"Created: {out}")


def generate_pdf():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # Try to register Chinese font
    font_paths = [
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
    ]
    font_name = "Helvetica"
    for fp in font_paths:
        if os.path.exists(fp):
            try:
                pdfmetrics.registerFont(TTFont("ChineseFont", fp))
                font_name = "ChineseFont"
                break
            except Exception:
                continue

    out = BASE / "等保合规自查报告.pdf"
    doc = SimpleDocTemplate(str(out), pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", parent=styles["Title"], fontName=font_name, fontSize=18, spaceAfter=12)
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName=font_name, fontSize=14, spaceAfter=8)
    body = ParagraphStyle("Body", parent=styles["Normal"], fontName=font_name, fontSize=10, leading=16, spaceAfter=6)

    story = []

    story.append(Paragraph("星云科技集团网络安全等级保护合规自查报告", title_style))
    story.append(Paragraph("报告年度：2025  |  报告类型：年度自查  |  密级：内部", body))
    story.append(Spacer(1, 12))

    story.append(Paragraph("一、报告概述", h1))
    story.append(Paragraph(
        "根据《信息安全等级保护管理办法》及 GB/T 22239-2019 要求，星云科技集团信息安全部"
        "于 2025 年 6 月至 7 月组织开展了年度网络安全等级保护合规自查。本次自查覆盖集团"
        "12 个等保三级系统、8 个等保二级系统，采用「文档审查 + 技术检测 + 人员访谈」"
        "相结合的方式，形成本报告。",
        body,
    ))

    story.append(Paragraph("二、自查范围", h1))
    scope_data = [
        ["序号", "系统名称", "等保级别", "备案号", "上次测评"],
        ["1", "星云OA办公系统", "三级", "33010613001-22001", "2024-11"],
        ["2", "星云ERP财务系统", "三级", "33010613001-22002", "2024-10"],
        ["3", "星云CRM客户系统", "三级", "33010613001-22003", "2025-01"],
        ["4", "星云eHR人事系统", "三级", "33010613001-22004", "2024-12"],
        ["5", "数据中台", "三级", "33010613001-22005", "2025-02"],
        ["6", "统一身份认证IAM", "三级", "33010613001-22006", "2024-09"],
        ["7", "安全运营中心SIEM", "三级", "33010613001-22007", "2025-03"],
        ["8", "互联网保险核心系统", "三级", "33010613001-22008", "2024-08"],
    ]
    t = Table(scope_data, colWidths=[1.2 * cm, 5 * cm, 2 * cm, 4.5 * cm, 2.5 * cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2F5496")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, -1), font_name),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(t)
    story.append(Spacer(1, 12))

    story.append(Paragraph("三、安全技术体系自查结果", h1))
    tech_data = [
        ["控制域", "检查项", "符合", "部分符合", "不符合"],
        ["安全物理环境", "机房访问控制、防火防水", "4", "1", "0"],
        ["安全通信网络", "边界防护、访问控制、入侵防范", "6", "2", "0"],
        ["安全区域边界", "WAF、防火墙策略、日志审计", "5", "1", "1"],
        ["安全计算环境", "身份鉴别、访问控制、安全审计", "12", "3", "1"],
        ["安全管理中心", "集中管控、日志汇聚、告警", "4", "2", "0"],
        ["安全管理制度", "制度体系、评审修订、培训", "8", "1", "0"],
        ["安全管理机构", "组织架构、人员配备、审批", "5", "0", "0"],
        ["安全管理人员", "录用离岗、考核培训、保密", "6", "1", "0"],
        ["安全建设管理", "定级备案、方案设计、测试验收", "7", "2", "0"],
        ["安全运维管理", "变更管理、备份恢复、应急预案", "9", "2", "1"],
    ]
    t2 = Table(tech_data, colWidths=[3.5 * cm, 5.5 * cm, 1.5 * cm, 2 * cm, 2 * cm])
    t2.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2F5496")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, -1), font_name),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
    ]))
    story.append(t2)
    story.append(Spacer(1, 12))

    story.append(Paragraph("四、主要差距与风险项", h1))
    gaps = [
        "1. 【中危】CRM 系统部分 API 接口未启用速率限制，存在暴力枚举风险；计划 2025-Q3 完成改造。",
        "2. 【中危】2 台遗留 Linux 服务器未安装 EDR 代理，已纳入资产清点专项，限期 2025-08-31 整改。",
        "3. 【低危】部分三级系统日志留存为 90 天，未达到制度要求的 180 天；SIEM 扩容项目已立项。",
        "4. 【低危】外包人员权限季度复核存在 2 个部门逾期；已发通报并纳入部门 KPI。",
        "5. 【信息】零信任网关 MFA 覆盖率 94%，目标 100%，剩余为测试专用账号。",
    ]
    for g in gaps:
        story.append(Paragraph(g, body))

    story.append(Paragraph("五、整改计划", h1))
    plan_data = [
        ["编号", "整改项", "责任部门", "完成时限", "状态"],
        ["R-01", "CRM API 限流改造", "研发效能部", "2025-09-30", "进行中"],
        ["R-02", "遗留服务器 EDR 覆盖", "IT运维中心", "2025-08-31", "进行中"],
        ["R-03", "日志留存扩容至180天", "信息安全部", "2025-10-31", "已立项"],
        ["R-04", "权限复核流程自动化", "IT运维中心", "2025-12-31", "规划中"],
        ["R-05", "零信任MFA 100%覆盖", "IT运维中心", "2025-08-15", "进行中"],
    ]
    t3 = Table(plan_data, colWidths=[1.2 * cm, 5 * cm, 3 * cm, 2.5 * cm, 2 * cm])
    t3.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2F5496")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, -1), font_name),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
    ]))
    story.append(t3)

    story.append(PageBreak())
    story.append(Paragraph("六、结论与建议", h1))
    story.append(Paragraph(
        "综合自查结果，星云科技集团等保三级系统整体符合率约 87.3%，较 2024 年度提升 4.2 个百分点。"
        "安全管理体系较为完善，主要差距集中在技术层面细粒度控制与运维自动化方面。",
        body,
    ))
    story.append(Paragraph("建议：", body))
    recs = [
        "持续投入零信任架构建设，提升远程办公与多云环境下的统一管控能力；",
        "推进权限复核与账号生命周期自动化，降低人工遗漏风险；",
        "加强开发安全（DevSecOps）能力，将安全检测嵌入 CI/CD 流水线；",
        "2025 年 Q4 邀请具有资质的测评机构开展正式等保测评。",
    ]
    for r in recs:
        story.append(Paragraph(f"• {r}", body))

    story.append(Spacer(1, 24))
    story.append(Paragraph("编制：信息安全部  |  审核：CISO  |  日期：2025年7月15日", body))

    doc.build(story)
    print(f"Created: {out}")


def generate_png():
    from PIL import Image, ImageDraw, ImageFont

    W, H = 1200, 900
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)

    # Try Chinese font
    font_paths = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
    ]
    title_font = None
    box_font = None
    small_font = None
    for fp in font_paths:
        if os.path.exists(fp):
            try:
                title_font = ImageFont.truetype(fp, 28)
                box_font = ImageFont.truetype(fp, 16)
                small_font = ImageFont.truetype(fp, 13)
                break
            except Exception:
                continue
    if title_font is None:
        title_font = box_font = small_font = ImageFont.load_default()

    draw.text((W // 2 - 200, 20), "安全事件应急响应流程图", fill="#1a1a2e", font=title_font)
    draw.text((W // 2 - 180, 55), "（P0/P1 级别 · 黄金四小时）", fill="#666", font=small_font)

    boxes = [
        (480, 100, 720, 160, "1. 发现与报告", "15分钟内初报SOC\n启动应急群"),
        (480, 200, 720, 260, "2. 研判与定级", "确认影响范围\nP0/P1/P2/P3"),
        (480, 300, 720, 360, "3. 遏制", "隔离主机/封禁IP\n保全日志镜像"),
        (480, 400, 720, 460, "4. 根除", "清除恶意代码\n修复漏洞"),
        (480, 500, 720, 560, "5. 恢复", "业务验证\n监控加强"),
        (480, 600, 720, 660, "6. 复盘", "7日内RCA报告\n更新预案"),
    ]

    for x1, y1, x2, y2, title, desc in boxes:
        draw.rounded_rectangle([x1, y1, x2, y2], radius=8, fill="#E8F4FD", outline="#2F5496", width=2)
        draw.text((x1 + 15, y1 + 8), title, fill="#2F5496", font=box_font)
        draw.text((x1 + 15, y1 + 32), desc, fill="#333", font=small_font)
        if y1 < 600:
            draw.line([(600, y2), (600, y2 + 40)], fill="#2F5496", width=2)
            draw.polygon([(590, y2 + 35), (610, y2 + 35), (600, y2 + 45)], fill="#2F5496")

    # Side branches
    branches = [
        (200, 300, "法务/公关\n评估对外通报"),
        (200, 400, "监管报告\n72小时内"),
        (900, 300, "业务Owner\n确认业务影响"),
        (900, 400, "IT运维\n隔离与恢复"),
    ]
    for x, y, text in branches:
        draw.rounded_rectangle([x, y, x + 180, y + 60], radius=6, fill="#FFF3E0", outline="#E65100", width=1)
        draw.text((x + 10, y + 10), text, fill="#E65100", font=small_font)
        draw.line([(x + 90 if x < 500 else x, y + 30), (480 if x < 500 else 720, 330 if y == 300 else 430)], fill="#E65100", width=1)

    # Legend
    draw.rounded_rectangle([50, 750, 1150, 870], radius=8, fill="#F5F5F5", outline="#CCC", width=1)
    draw.text((70, 760), "关键时限：", fill="#333", font=box_font)
    legends = [
        "P0：15分钟启动应急",
        "P1：30分钟响应",
        "遏制完成：≤4小时",
        "RCA报告：7个工作日",
        "通报模板：安全事件通报模板.txt",
    ]
    for i, leg in enumerate(legends):
        draw.text((70 + i * 210, 800), f"• {leg}", fill="#555", font=small_font)

    out = BASE / "应急响应流程图.png"
    img.save(out)
    print(f"Created: {out}")


if __name__ == "__main__":
    generate_docx()
    generate_xlsx()
    generate_pdf()
    generate_png()
    print("All binary documents generated.")
