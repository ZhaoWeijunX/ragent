#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate binary documents for onboarding knowledge base."""

import os
from pathlib import Path

BASE = Path(__file__).resolve().parent


def generate_docx():
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)

    title = doc.add_heading("星云科技集团岗位职责说明书", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("文件编号：XY-HR-JD-2025-001  |  版本：V2.0  |  密级：内部\n")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(80, 80, 80)
    meta.add_run("归口部门：人力资源部  |  适用对象：全体新入职员工参考")

    doc.add_paragraph()

    doc.add_heading("1. 文档说明", level=1)
    doc.add_paragraph(
        "本说明书汇总集团主要序列的典型岗位职责，供新员工了解岗位定位、"
        "核心职责与任职资格。具体岗位以 Offer 与劳动合同为准，"
        "部门经理可在通用模板基础上补充个性化内容。"
    )

    job_details = [
        {
            "title": "2. 研发工程师（后端）",
            "level": "P4-P6",
            "dept": "研发效能部 / AI平台部",
            "overview": "负责业务系统后端服务的设计、开发与维护，保障系统稳定性、性能与安全性。",
            "duties": [
                "参与需求评审，输出技术方案与工时评估；",
                "使用 Java/Go/Python 等语言完成模块开发与单元测试；",
                "编写 RESTful API，对接前端与第三方系统；",
                "参与 Code Review，遵循集团编码规范与安全要求；",
                "排查生产问题，参与 on-call 轮值（按部门安排）；",
                "撰写技术文档，参与知识库沉淀。",
            ],
            "requirements": [
                "本科及以上学历，计算机相关专业；",
                "2 年以上后端开发经验（校招应届生除外）；",
                "熟悉 Spring Boot、MySQL、Redis 等常用技术栈；",
                "了解分布式系统基础，有微服务经验者优先；",
                "良好的沟通协作能力与学习能力。",
            ],
            "kpis": ["需求按时交付率 ≥ 90%", "线上故障次数", "Code Review 参与度", "单元测试覆盖率"],
        },
        {
            "title": "3. 产品经理",
            "level": "P4-P6",
            "dept": "产品部",
            "overview": "负责产品规划、需求分析与项目推进，连接业务、研发与设计，交付用户价值。",
            "duties": [
                "进行市场调研与竞品分析，输出产品规划；",
                "撰写 PRD、用户故事，组织需求评审；",
                "跟进研发进度，协调资源，管理产品里程碑；",
                "与设计团队协作完成原型与交互方案；",
                "监控产品数据指标，持续迭代优化；",
                "收集用户反馈，维护需求池与版本路线图。",
            ],
            "requirements": [
                "本科及以上学历，理工科或商科背景；",
                "2 年以上互联网产品经验；",
                "熟练使用 Axure/Figma、Jira 等工具；",
                "具备数据分析能力，熟悉 SQL 者优先；",
                "优秀的逻辑思维与跨部门沟通能力。",
            ],
            "kpis": ["版本按时上线率", "核心功能使用率", "用户满意度 NPS", "需求变更控制"],
        },
        {
            "title": "4. HR 业务伙伴（HRBP）",
            "level": "P5-P7",
            "dept": "人力资源部",
            "overview": "作为业务部门的人力资源合作伙伴，推动人才发展、组织效能与员工关系管理。",
            "duties": [
                "深入了解所支持业务部门的战略与人才需求；",
                "参与招聘、入职、试用期管理与转正评估；",
                "推动绩效管理、晋升与调薪流程；",
                "处理员工关系问题，预防与化解劳动纠纷；",
                "组织部门培训需求调研与落地；",
                "协助推行集团 HR 政策与文化落地。",
            ],
            "requirements": [
                "本科及以上学历，人力资源或心理学相关专业；",
                "3 年以上 HRBP 或综合管理岗经验；",
                "熟悉劳动法规，持有 HR 证书者优先；",
                "出色的影响力与冲突处理能力；",
                "数据敏感度，能使用 eHR 系统产出人力分析。",
            ],
            "kpis": ["关键岗位到岗率", "试用期通过率", "员工敬业度", "离职率控制"],
        },
        {
            "title": "5. 销售经理",
            "level": "P4-P6",
            "dept": "销售部",
            "overview": "负责区域或行业客户的开拓、商机跟进与合同签署，完成销售目标。",
            "duties": [
                "开发目标客户，建立并保持客户关系；",
                "进行产品演示与方案讲解，推进商机转化；",
                "编制商务报价与投标文件，参与商务谈判；",
                "协调售前、交付团队完成签约后交接；",
                "维护 CRM 商机数据，定期提交销售预测；",
                "收集市场与竞品信息，反馈产品团队。",
            ],
            "requirements": [
                "大专及以上学历，市场营销或相关专业；",
                "2 年以上 B2B 软件或互联网销售经验；",
                "具备独立开拓客户与完成签约的能力；",
                "熟练使用 CRM 系统，有驾照者优先；",
                "抗压能力强，适应出差（约 30% 时间）。",
            ],
            "kpis": ["年度销售额达成率", "新客户签约数", "回款率", "商机转化率"],
        },
        {
            "title": "6. 运维工程师（SRE）",
            "level": "P4-P6",
            "dept": "IT运维中心",
            "overview": "负责生产环境基础设施与应用的稳定性保障、自动化运维与应急响应。",
            "duties": [
                "维护服务器、网络、容器平台等基础设施；",
                "建设监控告警体系，优化 SLI/SLO；",
                "参与故障应急响应与根因分析（RCA）；",
                "推进 IaC、CI/CD 与自动化运维脚本；",
                "执行变更管理，保障发布安全；",
                "参与等保与安全加固工作。",
            ],
            "requirements": [
                "本科及以上学历，计算机相关专业；",
                "2 年以上 Linux 运维或 SRE 经验；",
                "熟悉 Kubernetes、Prometheus、Ansible 等；",
                "了解网络与安全基础，有云平台经验；",
                "7×24 on-call 轮值意愿。",
            ],
            "kpis": ["系统可用性 SLA", "故障 MTTR", "变更成功率", "自动化覆盖率"],
        },
    ]

    for job in job_details:
        doc.add_heading(job["title"], level=1)
        doc.add_paragraph(f"职级范围：{job['level']}  |  所属部门：{job['dept']}")
        doc.add_heading("岗位概述", level=2)
        doc.add_paragraph(job["overview"])
        doc.add_heading("核心职责", level=2)
        for d in job["duties"]:
            doc.add_paragraph(d, style="List Bullet")
        doc.add_heading("任职资格", level=2)
        for r in job["requirements"]:
            doc.add_paragraph(r, style="List Bullet")
        doc.add_heading("试用期考核参考指标", level=2)
        for k in job["kpis"]:
            doc.add_paragraph(k, style="List Bullet")

    doc.add_heading("7. 通用行为要求（全员适用）", level=1)
    general = [
        "遵守《员工手册》及集团各项规章制度；",
        "保护公司商业秘密与客户数据，签署并履行保密协议；",
        "积极参与入职培训与部门周会，主动融入团队；",
        "在试用期内与直属经理保持至少每周一次 1:1 沟通；",
        "对岗位说明书有疑问时，向 HRBP 或直属经理咨询。",
    ]
    for g in general:
        doc.add_paragraph(g, style="List Bullet")

    doc.add_heading("8. 附则", level=1)
    doc.add_paragraph("本说明书由人力资源部负责维护，每年评审更新。")
    doc.add_paragraph("部门可在 eHR 上传个性化岗位说明作为本文件的补充。")

    out = BASE / "岗位职责说明书.docx"
    doc.save(out)
    print(f"Created: {out}")


def generate_xlsx():
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "培训课程表"

    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill("solid", fgColor="4472C4")
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    headers = [
        "课程编码", "课程名称", "类别", "对象", "形式",
        "时长(小时)", "讲师", "开课时间", "报名截止", "是否必修",
    ]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = thin

    courses = [
        ("ONB-001", "公司文化与价值观", "入职培训", "新员工", "线下", 2, "HR总监 陈静", "入职日当天", "—", "是"),
        ("ONB-002", "信息安全与保密", "入职培训", "新员工", "线上", 1.5, "信息安全部", "入职月内", "入职第2天", "是"),
        ("ONB-003", "考勤与假期制度", "入职培训", "新员工", "线上", 1, "HR共享中心", "入职月内", "入职第3天", "是"),
        ("ONB-004", "办公系统速成", "入职培训", "新员工", "线上", 1.5, "IT服务台", "入职月内", "入职第3天", "是"),
        ("ONB-005", "反骚扰与职场行为准则", "入职培训", "新员工", "线上", 1, "法务合规部", "入职月内", "入职第5天", "是"),
        ("GEN-001", "高效沟通与协作", "通用技能", "全员", "线下", 3, "外部讲师", "2025-09-15", "2025-09-10", "否"),
        ("GEN-002", "时间管理与目标拆解", "通用技能", "全员", "线上", 2, "内部讲师 李明", "2025-10-08", "2025-10-05", "否"),
        ("GEN-003", "结构化写作与汇报", "通用技能", "全员", "线下", 2.5, "外部讲师", "2025-11-12", "2025-11-08", "否"),
        ("DEV-001", "Java 进阶与性能优化", "专业技能", "研发", "线下", 8, "架构师 张工", "2025-09-20", "2025-09-15", "否"),
        ("DEV-002", "Spring Boot 3 实战", "专业技能", "研发", "线上", 6, "内部讲师", "2025-10-18", "2025-10-12", "否"),
        ("DEV-003", "AI/RAG 应用开发入门", "专业技能", "研发/产品", "线下", 4, "AI平台部", "2025-11-05", "2025-11-01", "否"),
        ("PRD-001", "需求分析与用户研究", "专业技能", "产品", "线下", 4, "产品总监", "2025-09-25", "2025-09-20", "否"),
        ("PRD-002", "数据驱动产品决策", "专业技能", "产品/运营", "线上", 3, "数据部", "2025-10-22", "2025-10-18", "否"),
        ("MGT-001", "新经理领导力工作坊", "管理者培训", "新任经理", "线下", 16, "外部机构", "2025-12-01", "2025-11-25", "否"),
        ("MGT-002", "绩效辅导与反馈技巧", "管理者培训", "管理者", "线下", 4, "HR L&D", "2025-09-08", "2025-09-03", "否"),
        ("MGT-003", "面试技巧与识人", "管理者培训", "面试官", "线下", 3, "HR招聘组", "2025-10-10", "2025-10-06", "否"),
        ("SAL-001", "大客户销售方法论", "专业技能", "销售", "线下", 6, "销售VP", "2025-09-18", "2025-09-12", "否"),
        ("OPS-001", "Kubernetes 运维实战", "专业技能", "运维", "线下", 8, "SRE团队", "2025-10-15", "2025-10-10", "否"),
        ("ONB-006", "星云产品全景介绍", "入职培训", "新员工", "线下", 2, "产品总监", "入职月内", "入职第7天", "否"),
        ("ONB-007", "Buddy 伙伴训练营", "入职培训", "Buddy", "线下", 2, "HR L&D", "每月首周五", "开课前3天", "否"),
    ]

    for row_idx, row_data in enumerate(courses, 2):
        for col_idx, val in enumerate(row_data, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.alignment = Alignment(vertical="center", wrap_text=True)
            cell.border = thin
            if col_idx == 10 and val == "是":
                cell.fill = PatternFill("solid", fgColor="E2EFDA")

    widths = [12, 24, 12, 12, 8, 10, 16, 14, 12, 10]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[chr(64 + i)].width = w

    ws2 = wb.create_sheet("培训类别说明")
    ws2["A1"] = "类别"
    ws2["B1"] = "说明"
    ws2["C1"] = "目标人群"
    for c in "ABC":
        ws2[f"{c}1"].font = header_font
        ws2[f"{c}1"].fill = header_fill

    cats = [
        ("入职培训", "帮助新员工快速融入公司", "入职 30 天内新员工"),
        ("通用技能", "跨岗位通用软技能", "全体员工"),
        ("专业技能", "岗位深度技能提升", "特定序列"),
        ("管理者培训", "团队管理与领导力", "经理及以上"),
    ]
    for i, row in enumerate(cats, 2):
        for j, val in enumerate(row, 1):
            ws2.cell(row=i, column=j, value=val)

    ws2.column_dimensions["A"].width = 14
    ws2.column_dimensions["B"].width = 30
    ws2.column_dimensions["C"].width = 22

    out = BASE / "培训课程表.xlsx"
    wb.save(out)
    print(f"Created: {out}")


def generate_pdf():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_paths = [
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
    ]
    font_name = "Helvetica"
    for fp in font_paths:
        if os.path.exists(fp):
            try:
                pdfmetrics.registerFont(TTFont("ChineseFont", fp))
                font_name = "ChineseFont"
                break
            except Exception:
                continue

    out = BASE / "员工手册.pdf"
    doc = SimpleDocTemplate(str(out), pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", parent=styles["Title"], fontName=font_name, fontSize=18, spaceAfter=12)
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName=font_name, fontSize=14, spaceAfter=8)
    body = ParagraphStyle("Body", parent=styles["Normal"], fontName=font_name, fontSize=10, leading=16, spaceAfter=6)

    story = []
    story.append(Paragraph("星云科技集团员工手册", title_style))
    story.append(Paragraph("2025 年版  |  内部资料  |  人力资源部编制", body))
    story.append(Spacer(1, 12))

    story.append(Paragraph("致员工的一封信", h1))
    story.append(Paragraph(
        "亲爱的星云同仁：欢迎加入星云科技集团。本手册是你了解公司制度、享受权利、"
        "履行义务的重要指引。请仔细阅读并遵守手册内容，如有疑问请联系 HRBP 或人力资源部。",
        body,
    ))

    story.append(Paragraph("第一章 公司简介", h1))
    story.append(Paragraph(
        "星云科技集团成立于 2015 年，总部位于杭州，在北京、上海、深圳设有分支机构。"
        "公司专注于企业级智能服务与 AI 应用，主要产品包括 Ragent 知识库平台、星云 OA、"
        "星云 CRM 及行业解决方案。截至 2025 年，集团员工超过 3,000 人。",
        body,
    ))

    story.append(Paragraph("第二章 聘用与试用期", h1))
    sections_2 = [
        "2.1 公司实行劳动合同制，劳动合同期限、试用期以合同约定为准。",
        "2.2 试用期一般为 3～6 个月，试用期内双方可依法解除劳动合同。",
        "2.3 试用期目标由直属经理与员工共同制定，期满前进行转正评估。",
        "2.4 转正须满足：试用期目标达成、无重大违纪、完成必修培训。",
        "2.5 员工应如实提供入职材料，如有虚假公司有权解除合同。",
    ]
    for s in sections_2:
        story.append(Paragraph(s, body))

    story.append(Paragraph("第三章 工作时间与考勤", h1))
    sections_3 = [
        "3.1 标准工作时间为 9:30～18:30，午休 12:00～13:30，每周工作 5 天。",
        "3.2 公司实行弹性工作制，可在 9:00～10:00 区间到岗，对应延迟下班。",
        "3.3 迟到早退、旷工按考勤制度处理；连续旷工 3 日视为严重违纪。",
        "3.4 加班须提前在 OA 申请并经审批，优先安排调休。",
        "3.5 考勤数据以企业微信或闸机记录为准，如有异议 5 日内申诉。",
    ]
    for s in sections_3:
        story.append(Paragraph(s, body))

    story.append(Paragraph("第四章 假期制度", h1))
    leave_data = [
        ["假期类型", "天数/规则", "审批要求"],
        ["年假", "按工龄：1-10年5天，10-20年10天，20年+15天", "直属经理"],
        ["病假", "凭医院证明，按当地政策发薪", "直属经理"],
        ["事假", "无薪，全年不超过15天", "直属经理"],
        ["婚假", "13天（按当地政策）", "HR备案"],
        ["产假", "按国家及地方生育政策", "HR备案"],
        ["陪产假", "15天", "HR备案"],
        ["丧假", "直系亲属3天", "直属经理"],
    ]
    t = Table(leave_data, colWidths=[3 * cm, 7 * cm, 4 * cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, -1), font_name),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
    ]))
    story.append(t)
    story.append(Spacer(1, 12))

    story.append(Paragraph("第五章 薪酬与福利", h1))
    sections_5 = [
        "5.1 薪酬由基本工资、岗位工资、绩效工资等组成，具体以 Offer 和劳动合同为准。",
        "5.2 发薪日为每月 10 日，通过银行转账发放。",
        "5.3 公司依法缴纳五险一金，并为员工购买商业补充医疗保险。",
        "5.4 福利包括：餐补、交通补贴（按岗位）、年度体检、节日礼品、团建经费等。",
        "5.5 薪酬属于保密信息，员工不得打听、议论他人薪酬。",
    ]
    for s in sections_5:
        story.append(Paragraph(s, body))

    story.append(PageBreak())
    story.append(Paragraph("第六章 行为准则与纪律", h1))
    sections_6 = [
        "6.1 员工应诚实守信、爱岗敬业、团结协作。",
        "6.2 禁止收受贿赂、回扣、礼品（超过 200 元须报备）。",
        "6.3 禁止在工作时间从事兼职或损害公司利益的活动。",
        "6.4 禁止泄露公司商业秘密、客户数据及员工个人信息。",
        "6.5 禁止性骚扰、歧视、暴力等违反职场行为准则的行为。",
        "6.6 违反纪律的，按情节给予警告、记过、降职、解除合同等处理。",
    ]
    for s in sections_6:
        story.append(Paragraph(s, body))

    story.append(Paragraph("第七章 知识产权", h1))
    story.append(Paragraph(
        "员工在职期间因履行职务或主要利用公司资源完成的发明创造、软件著作权、"
        "技术秘密及其他知识产权归公司所有。离职后仍须对在职期间知悉的商业秘密承担保密义务。",
        body,
    ))

    story.append(Paragraph("第八章 离职管理", h1))
    sections_8 = [
        "8.1 员工辞职须提前 30 日（试用期 3 日）书面通知。",
        "8.2 离职须完成工作交接、归还公司财物、注销账号。",
        "8.3 公司可在员工严重违纪等法定情形下解除劳动合同。",
        "8.4 离职证明在办结手续后 15 日内出具。",
    ]
    for s in sections_8:
        story.append(Paragraph(s, body))

    story.append(Spacer(1, 24))
    story.append(Paragraph("本人已阅读并理解本手册内容，承诺遵守公司各项规章制度。", body))
    story.append(Spacer(1, 12))
    story.append(Paragraph("员工签字：________________    日期：____年__月__日", body))
    story.append(Spacer(1, 24))
    story.append(Paragraph("人力资源部  |  2025 年 1 月 1 日生效", body))

    doc.build(story)
    print(f"Created: {out}")


def generate_png():
    from PIL import Image, ImageDraw, ImageFont

    W, H = 1200, 800
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)

    font_paths = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
    ]
    title_font = box_font = small_font = None
    for fp in font_paths:
        if os.path.exists(fp):
            try:
                title_font = ImageFont.truetype(fp, 26)
                box_font = ImageFont.truetype(fp, 14)
                small_font = ImageFont.truetype(fp, 12)
                break
            except Exception:
                continue
    if title_font is None:
        title_font = box_font = small_font = ImageFont.load_default()

    draw.text((W // 2 - 180, 15), "星云科技集团组织架构图", fill="#1a1a2e", font=title_font)
    draw.text((W // 2 - 120, 48), "（2025 年 1 月版）", fill="#666", font=small_font)

    # Top: Board / CEO
    draw.rounded_rectangle([480, 80, 720, 130], radius=6, fill="#4472C4", outline="#2F5496", width=2)
    draw.text((530, 95), "董事会 / CEO", fill="white", font=box_font)

    draw.line([(600, 130), (600, 160)], fill="#2F5496", width=2)

    # C-level
    cx_boxes = [
        (120, 160, 280, 210, "CTO\n技术委员会"),
        (340, 160, 500, 210, "CPO\n产品委员会"),
        (520, 160, 680, 210, "COO\n运营委员会"),
        (700, 160, 860, 210, "CFO\n财务委员会"),
        (920, 160, 1080, 210, "CHRO\n人力资源"),
    ]
    for x1, y1, x2, y2, text in cx_boxes:
        draw.rounded_rectangle([x1, y1, x2, y2], radius=6, fill="#E8F4FD", outline="#4472C4", width=1)
        lines = text.split("\n")
        draw.text((x1 + 20, y1 + 12), lines[0], fill="#2F5496", font=box_font)
        draw.text((x1 + 20, y1 + 32), lines[1], fill="#555", font=small_font)
        draw.line([(600, 160), ((x1 + x2) // 2, y1)], fill="#2F5496", width=1)

    draw.line([(600, 210), (600, 240)], fill="#2F5496", width=2)

    # Business units
    bu_boxes = [
        (60, 240, 240, 300, "研发效能部\n(4F-7F)"),
        (260, 240, 440, 300, "AI平台部\n(5F)"),
        (460, 240, 640, 300, "产品部\n(8F)"),
        (660, 240, 840, 300, "销售部\n(10F)"),
        (860, 240, 1040, 300, "市场部\n(9F)"),
        (160, 320, 340, 380, "客户服务部\n(B座3F)"),
        (380, 320, 560, 380, "数据部\n(中台)"),
        (580, 320, 760, 380, "IT运维中心\n(7F)"),
        (780, 320, 960, 380, "信息安全部\n(3F)"),
    ]
    for x1, y1, x2, y2, text in bu_boxes:
        draw.rounded_rectangle([x1, y1, x2, y2], radius=6, fill="#F5F5F5", outline="#999", width=1)
        lines = text.split("\n")
        draw.text((x1 + 15, y1 + 12), lines[0], fill="#333", font=box_font)
        draw.text((x1 + 15, y1 + 32), lines[1], fill="#888", font=small_font)

    # Support functions
    draw.rounded_rectangle([400, 420, 800, 470], radius=6, fill="#FFF3E0", outline="#E65100", width=1)
    draw.text((430, 435), "共享服务：人力资源 | 财务共享 | 行政 | 法务合规 | 采购", fill="#E65100", font=small_font)

    # Legend
    draw.rounded_rectangle([50, 520, 1150, 760], radius=8, fill="#FAFAFA", outline="#CCC", width=1)
    draw.text((70, 535), "新员工须知：", fill="#333", font=box_font)
    tips = [
        "1. 入职手续在 A座3F HR服务中心办理",
        "2. 研发默认工位在 4F-7F，具体见 Offer",
        "3. 直属经理是绩效考核与日常审批的第一负责人",
        "4. HRBP 负责所支持事业部的人事政策咨询",
        "5. 详细楼层与门禁见《办公地点与门禁卡.csv》",
    ]
    for i, tip in enumerate(tips):
        draw.text((70, 570 + i * 32), tip, fill="#555", font=small_font)

    out = BASE / "组织架构图.png"
    img.save(out)
    print(f"Created: {out}")


if __name__ == "__main__":
    generate_docx()
    generate_xlsx()
    generate_pdf()
    generate_png()
    print("All binary documents generated.")
